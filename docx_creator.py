import docx
from docx import Document
from datetime import datetime
import os
import glob
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def create_docx(topic, blog_post, images, yoast_seo_data, tags, facebook_posts):
    """
    Create a DOCX file with the given content and save it in a directory named with the current date.
    Automatically increments the version number to avoid overwriting existing files.

    :param topic: The topic of the blog post.
    :param blog_post: The content of the blog post.
    :param images: A list of image file paths.
    :param yoast_seo_data: The SEO data.
    :param tags: The list of tags.
    """
    # Create a new document
    doc = Document()

    # Add content to the document
    doc.add_heading(topic, 0)
    doc.add_paragraph(blog_post)
    doc.add_paragraph(facebook_posts)
    for image in images:
        doc.add_picture(image, width=docx.shared.Inches(5))

    # Add SEO data and tags
    doc.add_heading('SEO Data', level=1)
    doc.add_paragraph(f"Focus Keyphrase: {yoast_seo_data['focus_keyphrase']}")
    doc.add_paragraph(f"SEO Title: {yoast_seo_data['seo_title']}")
    doc.add_paragraph(f"Meta Description: {yoast_seo_data['meta_description']}")
    doc.add_paragraph("Tags: " + ", ".join(tags))

    # Add Facebook posts
    doc.add_heading('Facebook Posts', level=2)
    for post in facebook_posts:
        paragraph = doc.add_paragraph(post)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Adjust alignment as needed


    # Directory and file handling
    today = datetime.now().strftime("%Y-%m-%d")
    save_directory = os.path.join("generated_docs", today)
    os.makedirs(save_directory, exist_ok=True)

    # Determine the next version number
    existing_files = glob.glob(os.path.join(save_directory, f"{topic}_v*_{today}.docx"))
    existing_versions = [int(file.split('_v')[1].split('_')[0]) for file in existing_files]
    next_version = max(existing_versions, default=0) + 1

    # Construct file name with topic, next version number, and current date
    filename = f"{topic}_v{next_version}_{today}.docx"
    file_path = os.path.join(save_directory, filename)

    # Save the document
    doc.save(file_path)
    print(f"Document saved as {file_path}")
